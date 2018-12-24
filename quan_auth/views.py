from .models import *

from quan_badge.models import *
from quan_event.models import *

from quan_account.views import *
from quan_avatar.views import AvatarSnap
from quan_event.views import EventSnap
from quan_message.views import MessageSnap
from quan_message.views import PublicMessageAction
from quan_share.views import ShareSnap, ShareDicts
from quan_ua.views import *

class AuthSnap :
	def qrank_calc(cid) :
		'计算社团星级（0.0-5.0）'
		'''
			分值			0	1	2	3	4	5
						========================
			每月活动次数	0	1	2	4	6	8
			活跃成员比例	15%	30%	45%	60%	75%	90%
			月底提交记录	False				True
			成员总数		0	4	8	12	16	20

			附加分*0.4：
			关注人数		0	8	16	24	32	36

			计算方法：
				记总和，除以总数4，向上取整。
				>5 => 5
		'''
		month_event = EventSnap.event_month_find_by_cid('club', cid).count()
		member_sum = ClubSnap.member_total(cid)
		active_ratio = AuthSnap.active_ratio(cid)
		submit_record = 1										# 暂时不考虑
		follower_total = ClubSnap.follower_total(cid)
		# 计算
		sum_score = 0.0
		sum_score += (0, 1, 2, 2.5, 3, 3.5, 4, 4.5, 5)[min(month_event, 8)]
																# 每月活动数量
		sum_score += min(max(active_ratio / 0.15 - 1, 0), 5)	# 活跃成员
		sum_score += 5 * int(submit_record)						# 是否提交
		sum_score += min(member_sum / 4, 5)						# 成员总数
		addition_score = min(follower_total / 8, 5)				# 附加分 关注人数
		# 求和
		sum_score /= 4
		sum_score += 0.4 * addition_score
		return min(sum_score, 5)

	def mrank_calc(uid, cid) :
		'计算社员等级（1.0-5.0）'
		'''
			分值								0	1	2	3	4	5
											========================
			参与率		ratio_join			15%	30%	45%	60%	75%	90%
			加入时间		month_join			0mo	1mo	2mo	3mo	4mo	6mo	
			参与小时数	attendence_hour		2	4	6	8	10	12
		'''
		#得到数据
		ratio_join = AuthSnap.ratio_join(uid, cid)
		month_join = ClubSnap.attendence_total(uid, cid)
		attendence = EventSnap.event_attendence_by_uid(uid, cid, True)[0]
		#计算
		return (
			  min(ratio_join / 0.15, 5)
			+ [0, 1, 2, 3, 4, 4.5, 5][min(int(month_join.days / 30), 6)]
			+ min(max(attendence / timedelta(0, 3600) - 1, 0), 5)
		) / 3

	def ratio_join(uid, cid) :
		'获得用户在社团中一个月的参与率(0~1)'
		join, total = EventSnap.event_attendence_by_uid(uid, cid, True)
		try :
			return join / total
		except ZeroDivisionError :
			return 0

	def active_check(uid, cid) :
		'返回社员是否活跃（60 percent 时间以上）'
		return AuthSnap.ratio_join(uid, cid) > 0.6
		
	def active_total(cid) :
		'返回社团的活跃成员总数'
		from quan_account.views import AccountSnap
		active_num = 0
		for i in ClubSnap.member_all(cid) :
			active_num += int(AuthSnap.active_check(i, cid))
		return active_num

	def notice_read(uid) :
		'对于读取通知的数据库操作，成功返回True'
		if not uid :
			return False
		AuthInfo(account_id=uid, relation='read-notice').save()
		return True

	def notice_check(uid, note) :
		'是否已经读取通知，note为通知数据库表'
		if not note :
			return True
		if not uid :
			return False
		return AuthInfo.objects.filter(relation='read-notice', account_id=uid, 
							time_create__gte=note['data'].time_create).exists()

	def amercement_add(cid, data) :
		'添加坏行为记录，成功返回ID，失败返回None'
		if not cid :
			return None
		qry = AuthInfo(relation='amercement', account_id=cid, data=data)
		qry.save()
		return qry.id

	def amercement_del(amid) :
		'删除坏行为记录，返回True'
		return bool(AuthInfo.objects.filter(relation='amercement', id=amid) \
								.update(relation='amercement-history'))

	def amercement_num(cid) :
		'返回社团的违反次数'
		qry = AuthInfo.objects.filter(relation='amercement', account_id=cid)
		return qry.count()

	def active_ratio(cid) :
		'社团活跃成员比例'
		member_sum = ClubSnap.member_total(cid)
		if member_sum :
			active_sum = AuthSnap.active_total(cid)
			if active_sum :
				return active_sum / member_sum
		return 0

	def join_ratio(cid) :
		event_list = EventSnap.event_month_find_by_cid('club', cid)
		active_rate = 0
		for event in event_list :
			active_rate += EventSnap.signup_total(event.id)
		if event_list.count() :
			return active_rate / event_list.count()
		else :
			return 0

	def category_list(uid) :
		'一个用户可以管理联盟的代号列表'
		cuid = ClubSnap.cid_find_by_sname('club-union')
		qry = BadgeRelation.objects.filter(send_id=cuid, send_relation='club', 
				status=0, recv_id=uid, recv_relation='user').values_list \
				('reason', flat=True)
		category_list = set(map(lambda x: x[0], ClubSnap.category_list()))
		return list(filter(lambda x: x in category_list, qry))

	def star_calc(cid, event_qry=None, member_makeup=None, share_qry=None, 
					return_score=False) :
		'''
			计算星级
				本学期活动数量	* 10 		|	8	6	4	3	2
				发布分享数量	* 10		|	4	3	2	1	0
				成员数量		* 4			|	15	12	9	6	3
				成员年级数量	* 20		|	5	4	3	2	1
				类加后，使用以下评分标准
				五星: >= 80				|	max: 280
				四星: >= 70				|	max: 218
				三星: >= 60				|	max: 156
				二星: >= 50				|	max: 104
				一星: >= 40				|	max: 52
		'''
		#准备
		if event_qry == None :
			event_qry = EventSnap.event_semester_find_by_cid('club', cid).order_by \
				('time_set')
		if member_makeup == None :
			member_makeup = [0, 0, 0, 0, 0, 0, 0, 0]
			for i in ClubSnap.grade_list_by_cid(cid) :
				member_makeup[i] += 1
		if share_qry == None :
			start = EventSnap.semester_start()
			share_qry = ShareSnap.share_find_by_aid('club', cid)
			share_qry = share_qry.filter(time_create__gte=start)
		#开始计算
		member_sum = sum(member_makeup[:-1])
		member_grade_num = 0
		for i in member_makeup :
			member_grade_num += int(bool(i))
		event_num = event_qry.count()
		share_num = share_qry.count()
		#开始累加
		score = event_num * 10 + share_num * 10 + \
				member_sum * 4 + member_grade_num * 20
		if return_score :
			return score
		#返回
		if score >= 80 :
			return 5
		if score >= 70 :
			return 4
		if score >= 60 :
			return 3
		if score >= 50 :
			return 2
		if score >= 40 :
			return 1
		return 0

	def form_fetch(cid) :
		'得到社团最后的申请表'
		from quan_share.models import ShareInfo, ShareAttach
		share_qry = ShareInfo.objects.filter \
			(status=0, account_id=cid, relation='club', category='form')
		uuid_list = []
		for i in share_qry.order_by('-time_update') :
			uuid_list.append(i.attach_uuid)
		qry = ShareAttach.objects.filter(attach_uuid__in=uuid_list, 
				category='file').order_by('-time_create')
		if qry :
			sname = AccountSnap.sname_find_by_aid(qry[0].relation, 
												qry[0].account_id)
			spath = os.path.join('files', sname, qry[0].name_new)
			file_name = settings.MEDIA_ROOT + spath
			return file_name
		else :
			return None

	def apply_statistics() :
		'计算每个星级的当前申报数量，暂时没有被非shell引用'
		h = ['preparing', 'waiting', 'denied', 'droped', 'broken', 'granted']
		t = {}
		for i in range(len(h)) :
			t[h[i]] = i
		ans = []
		for i in range(6) :
			ans.append([0] * len(h))
		for i in AuthInfo.objects.filter(relation='performance') :
			ans[int(i.stars)][t[i.data]] += 1
		hh = ['正在填写', '等待审核', '未通过审核', '材料不足', '材料损坏', '审核通过']
		print('star', *hh, sep='\t')
		for i in range(6) :
			print(i, *ans[i], sep='\t')
		return (h, ans)

	def funds_document_generate(club, head, extra, real_name, raw_name, qry) :
		from docx import Document
		template = os.path.join(settings.STATIC_ROOT, 'funds/funds_model.docx')
		docx = Document(template)
		# { 'funds': '社团发展基金', 'borrow': '社团借款', }[extra[1]], 
		head_account = UserSnap.account_find_by_uid(head.id)
		translate = {
			'あ': '20' + str(qry.stars), 
			'い': club.full_name, 
			'う': head.first_name, 
			'え': UserSnap.grade_translate(head_account.grade), 
			'お': head_account.phone, 
			'か': str(extra[2]), 
			'き': qry.time_create.strftime('%Y年%m月%d日'), 
			'く': extra[0], 
			'け': extra[3], 
			'こ': extra[4], 
						# extra[5] 是邮箱
			'さ': '', 	# extra[6]
		}
		if len(extra) > 6 :
			translate['さ'] = extra[6]
		for i in docx.paragraphs :
			if 'あ' in i.text :
				i.text = i.text.replace('あ', '20' + str(qry.stars))
		for i in docx.tables :
			for j in i.rows :
				for k in j.cells :
					if k.text in translate :
						k.text = translate[k.text]
		docx.save(real_name)

	def get_rubric(category='global') :
		'得到某个联盟或总体 (global) 的 rubric ，失败返回 []'
		qry = AuthInfo.objects.filter(relation='performance-rubric', 
				data=category, time_update__gte=EventSnap.semester_start())
		if qry :
			return json.loads(qry[0].extra)
		else :
			return []

class AuthDicts :
	def club_list(qry) :
		'为/auth/edit/提供社团列表，返回列表'
		club_list = []
		for i in qry :
			club_list.append({
				'data':				i, 
				'crank':			int(AuthSnap.qrank_calc(i.id)), 
				'attend_ratio':		int(AuthSnap.join_ratio(i.id) * 100), 
				'member_num':		ClubSnap.member_total(i.id), 
				'amercement_num':	AuthSnap.amercement_num(i.id),
			})
		return club_list

	def auth_apply(is_auth) :
		'is_auth -> no preparing'
		dict_render = {'record_list': []}
		qry = AuthInfo.objects.filter(relation="performance", 
			time_update__gte=EventSnap.semester_start()).order_by('-time_update')
		label_dict = {
			'preparing': '填写中', 
			'waiting': '等待中', 
			'granted': '审核通过', 
			'denied': '审核不通过', 
			'droped': '材料不足', 
			'broken': '附件损坏', 
		}
		label_type_dict = {
			'preparing': 'default', 
			'waiting': 'info', 
			'granted': 'success', 
			'denied': 'danger', 
			'droped': 'warning', 
			'broken': 'warning', 
		}
		category_list = ClubSnap.category_list()
		dict_render['category_dict'] = category_list
		category_dict = dict(category_list)
		for i in qry :
			if i.data == 'preparing' and is_auth :
				continue
			club = ClubSnap.club_find_by_cid(i.account_id)
			if not club :
				continue	# 数据错误
			date, time = ChineseSnap.datetime_simp(i.time_update)
			data = {
				'data': i, 
				'club': club, 
				'date': date + time, 
				'zhcn_label': label_dict[i.data], 
				'zhcn_category': category_dict[club.category], 
				'label_type': label_type_dict[i.data], 
				'amercement': AuthSnap.amercement_num(i.account_id), 
			}
			dict_render['record_list'].append(data)
		return dict_render

	def page_dict(request, all_list, elem_num, dict_render, index=None) :
		'''
			将一个列表中的某一段取出
			参数
				request: 通过 GET['index'] 参数得到页码
				all_list: 完整列表。如果传递 QuerySet 会有更高的效率
					例如， User.objects.filter() 是 QuerySet
				elem_num: 每页的元素数量
				dict_render: 将渲染模板的参数更新到 dict_render 中
		'''
		if type(all_list) == QuerySet :
			page_len = max(1, (all_list.count() - 1) // elem_num + 1)
		else :
			page_len = max(1, (len(all_list) - 1) // elem_num + 1)

		try :
			if index == None :
				index = int(request.GET.get('index', '1'))
		except ValueError :		# 非法的 index
			index = 1
		if index <= 1 :			# 列表开头
			index = 1
		else :
			dict_render['page_pre'] = index - 1
		if index >= page_len :	# 列表结尾
			index = page_len
		else :
			dict_render['page_next'] = index + 1
		
		dict_render['page_len'] = page_len
		dict_render['page_list'] = range(max(1, index - 4), 
										min(page_len + 1, index + 5))
		dict_render['page_index'] = index
		return all_list[(index - 1) * elem_num: index * elem_num]

	def funds_status_translate() :
		return {
			'preparing': '正在填写', # 可选
			'deleted': '被社长删除', # 可选
			'submitted': '已提交', 
			'head-wait': '已审核待社长查看', 
			'head-agree': '已审核社长同意', 
			'head-deny': '已审核社长拒绝', 
			'granted': '申请通过', 
			'rejected': '申请失败', 
			'droped': '材料缺失', 
		}

	def funds_list(qry, need_club=True) :
		'''
			funds 中的列表渲染
			need_club: 是否返回社团名称
		'''
		status_translate = AuthDicts.funds_status_translate()
		funds_list = []
		for i in qry :
			if i.data == 'deleted' :
				continue
			extra = json.loads(i.extra)
			funds_list.append({
				'data': i, 
				'type': {'funds': '社团发展基金', 'borrow': '社团借款', }[extra[1]], 
				'status': status_translate[i.data], 
				'extra': extra, 
			})
			if need_club :
				funds_list[-1]['club'] = ClubSnap.club_find_by_cid(i.account_id)
		return funds_list

	def funds_form() :
		'返回表格'
		return [
#			('标题', '填写标题', ), 
#			('类型', '选择类型', ), 
			# 序号, label, placeholder, 高度
			(1, '申请金额', '输入整数人民币金额', 1, ), # [2]
			(2, '学期活动规划', '学期活动规划', 15, ), # [3]
			(3, '经费使用说明', '经费使用说明', 15, ), # [4]
			(4, '社长邮箱', '请填写常用邮箱并在近期注意查看', 1, ), # [5]
		]

class AuthViews :
	@vary_on_cookie
	@UserAuthority.club_union_check
	def auth_manage(request) :
		dict_render = UserAgentSnap.record(request)
		if request.method == 'GET' :
			qry = ClubSnap.club_all_find().order_by('-time_create')
			sliced_qry = AuthDicts.page_dict(request, qry, 10, dict_render)
			dict_render['club_list'] = AuthDicts.club_list(sliced_qry)
			dict_render['title'] = ' - 管理社团'
			return Snap.render('auth_manage.html', dict_render)
		else :
			action_type = request.POST.get('type', '')
			cid = int(request.POST.get('cid', '0'))
			msg = request.POST.get('msg', '')
			if action_type == 'remove' :
				raise Snap.error('由于此功能太危险，目前已经被禁用')
				ClubSnap.club_remove(cid)
				return Snap.success(request, '成功删除社团，请刷新页面')
			elif action_type == 'amerce' :
				amid = AuthSnap.amercement_add(cid, msg)
				if not amid :
					raise Snap.error('没有找到相关社团')
				hid = ClubSnap.head_id_find_by_cid(cid)
				csname = ClubSnap.sname_find_by_cid(cid)
			#	unid = ClubSnap.cid_find_by_sname('club-union')
			#	MessageSnap.qmsg_push('club-amerce', amid, 'user', hid, 
			#			'[处罚] 名下的%s被社联给予处罚' % csname, 
			#			'请遵守社团联合会的相关规章制度。', 'club', unid)
				return Snap.success(request, '成功对涉事社团进行处罚，请刷新页面')
			else :
				raise Snap.error('似乎在对我做一些不好的事情')
			
	@vary_on_cookie
	def auth_apply(request) :
		dict_render = UserAgentSnap.record(request)
		is_auth = ClubSnap.club_union_check(request)
		if request.method == 'GET' :
			dict_render.update(AuthDicts.auth_apply(is_auth))
			dict_render['is_auth'] = is_auth
			uid = UserSnap.uid_find_by_request(request)
			category_list = AuthSnap.category_list(uid)
			category_dict = dict(ClubSnap.category_list())
			dict_render['category_manage'] = \
					list(map(lambda x: (x, category_dict[x]), category_list))
			if is_auth :
				dict_render['category_manage'].insert(0, ('global', t_('综合')))
			dict_render['title'] = ' - 星级评价'
			return Snap.render('auth_apply.html', dict_render)
		else :
			if not is_auth :
				raise Http403("如果为社联成员，请检查登录情况。")
			try :
				pid = int(request.POST.get('pid', ''))	# performance id
			except Exception :
				raise Snap.error('表单信息有误')
			result = request.POST.get('result', '')
			qry = AuthInfo.objects.filter(id=pid, relation='performance')
			if qry.count() != 1 or not result in \
				('granted', 'denied', 'droped', 'broken') :
				raise Snap.error('表单信息有误')
			if qry[0].time_update < EventSnap.semester_start() :
				raise Snap.error('表单信息有误')
		#	if qry[0].category != 'record-waiting' :
		#		raise Snap.error('请求早已处理')
			qry.update(data=result)
			#fname = ClubSnap.fname_find_by_cid(qry[0].account_id)
			#hid = ClubSnap.head_id_find_by_cid(qry[0].account_id)
			#unid = ClubSnap.cid_find_by_sname('club-union')
			if result == 'granted' :
				PublicMessageAction.save \
					(request, 'empty', 'granted', qry[0].account_id, 'club', 
					ClubSnap.cid_find_by_sname('club-union'), 'club')
			else :
				PublicMessageAction.save \
					(request, 'empty', 'fine', qry[0].account_id, 'club', 
					ClubSnap.cid_find_by_sname('club-union'), 'club')
			return Snap.success(request, '成功执行请求。')

	@vary_on_cookie
	def auth_apply_rubric(request, category) :
		dict_render = UserAgentSnap.record(request)
		# 检测权限
		uid = UserSnap.uid_find_by_request(request)
		if category == 'global' :
			if not ClubSnap.club_union_check(request) :
				raise Http403('您没有社联管理权限')
		else :
			if category not in AuthSnap.category_list(uid) :
				raise Http403('您没有联盟管理权限')
		# 查询是否已存在
		qry = AuthInfo.objects.filter(relation='performance-rubric', 
					data=category, time_update__gte=EventSnap.semester_start())
		if qry :
			qry = qry[0]
		if request.method == "GET" :
			dict_render['category'] = category
			if qry :
				data = json.loads(qry.extra)
				dict_render['form_data'] = '\n'.join(map(
							lambda x: '\t'.join([x[0]] + list(map(str, x[1]))), 
							data
						))
				dict_render['best_score'] = sum(map(lambda x: x[1][-1], data))
			dict_render['title'] = ' - 编辑星级评分表'
			return Snap.render('auth_apply_rubric.html', dict_render)
		data = request.POST.get('data')
		if not data :
			raise Snap.error('表单内容不全')
		a = map(lambda x: x.split(), filter(bool, data.split('\n')))
		answer = []
		for index, i in enumerate(a) :
			row = []
			for jndex, j in enumerate(i[1:]) :
				try :
					row.append(int(j))
				except ValueError :
					try :
						float_j = float(j)
						if float_j * 2 % 1 :
							raise Snap.error(t_('第%d行%d列：只能精确到0.5') % 
													(index + 1, jndex + 2))
						int_j = int(float_j)
						if int_j == float_j :
							row.append(int_j)
						else :
							row.append(float_j)
					except ValueError :
						raise Snap.error(t_('第%d行%d列数据有误') % 
													(index + 1, jndex + 2))
			if not row :
				raise Snap.error(t_('第%d行没有数据') % (index + 1))
			if len(row) != len(set(map(float, row))) :
				raise Snap.error(t_('第%d行有重复数据') % (index + 1))
			answer.append((i[0], sorted(row)))
		if qry :
			qry.extra = json.dumps(answer)
			qry.account_id=uid
			qry.save()
			return Snap.success(request, '更改成功', 
				{ 'redirect': '/auth/apply/rubric/%s/' % category })
		else :
			AuthInfo(extra=json.dumps(answer), relation='performance-rubric', 
						data=category, account_id=uid).save()
			return Snap.success(request, '保存成功', 
				{ 'redirect': '/auth/apply/rubric/%s/' % category })

	@vary_on_cookie
	def room_apply(request) :
		dict_render = UserAgentSnap.record(request)
		dict_render['title'] = ' - 活动空间申请'
		return Snap.render('room_apply.html', dict_render)

	@vary_on_cookie
	def empty_view(request, sname) :
		dict_render = UserAgentSnap.record(request)
		cid = ClubSnap.cid_find_by_sname(sname)
		is_head = AccountSnap.aid_verify(request, 'club', cid)
		is_union = AccountSnap.aid_verify(request, 'club', 
					ClubSnap.cid_find_by_sname('club-union'), True, True)
		category_list = AuthSnap.category_list \
			(UserSnap.uid_find_by_request(request))
		club = ClubSnap.club_find_by_cid(cid)
		if not is_head and not is_union and not category_list :
			raise Http403 \
				('根据社联方面的要求，只有社长和社联成员才能查看此页面。请确认登录状态')
		if not club :
			raise Http404('数据错误，请联系HCC社团')
		is_category_admin = club.category in category_list
		performance_qry = AuthInfo.objects.filter(relation='performance', 
					account_id=cid, time_update__gte=EventSnap.semester_start())
		# get performance
		if performance_qry :
			performance = performance_qry[0]
		else :
			performance = AuthInfo \
				(relation='performance', account_id=cid, data='preparing')
			performance.save()
		if not is_head and not is_union and performance.data == 'preparing' :
			raise Http403('表格尚未公开')
		if request.method == 'GET' :
			# 处理提交的表单
			dict_render['performance'] = performance
			dict_render['changeable'] = performance.data in \
				('waiting', 'preparing')
			dict_render['event_location'] = request.GET.get('event_location', '')
			dict_render['event_subject'] = request.GET.get('event_subject', '')
			dict_render['time_set'] = request.GET.get('time_set', '')
			dict_render['shareSubject'] = request.GET.get('shareSubject', '')
			impact_val = []
			for i in range(5) :
				impact_val.append(request.GET.get('impact%d' % i, ''))
			if not any(impact_val) :
				try :
					impact_val = json.loads(performance.extra)
				except ValueError :
					pass
			# begin rendering
			impact_form = ['社团活动反思总结', '与其他社团合作活动', 
							'承办学校活动项目', '社团成果', '加分项申请']
			impact_placeholder = [
				'社团本学期的大致活动内容，例如日常开会、日常研讨等。'
				'若有固定时间也请进行补充说明', 
				'合作活动的名称、时间以及合作社团的名称，外加合作社团社长的联系方式与姓名。'
				'需要上传可视化资料进行证明，例如照片视频等（在分享中上传）', 
				'学校活动的名称、时间、50字简述。'
				'需要上传可视化资料进行证明，例如照片视频等（在分享中上传）', 
				'本学期该社团具体都达成了什么成就，如果出杂志刊物的就列举出刊物和刊期', 
				'根据评审方法上的加分项内容进行填写', 
				'暂定是这些，谢谢', 
			]
			# 以上为 2015-2016 下学期的版本
			impact_form = ['举办过的社团内部活动', '参与或举办的校级活动', 
						'举办的独立社团展示', '本学期社团活动总结与反思', '加分项申请']
			impact_placeholder = [
				'写明次数、时间地点、活动形式等', 
				'活动名称和活动形式', 
				'请参照核心评分标准填写', 
				'字数大约为100字', 
				'照片、视频、文档等证明资料为星级评价的重要标准，请在上方的分享栏中提交', 
			]
			dict_render['impact'] = []
			for i in range(5) :
				dict_render['impact'].append((i, impact_form[i], impact_val[i], 
												impact_placeholder[i]))
			semester_start = EventSnap.semester_start()
			dict_render['is_head'] = is_head
			dict_render['is_union'] = is_union
			dict_render['is_category_admin'] = is_category_admin
			head = ClubSnap.head_find_by_cid(cid)
			dict_render['head'] = head
			head_account = UserSnap.account_find_by_uid(head.id)
			dict_render['head_phone'] = head_account.phone
			dict_render['avatar'] = AvatarSnap.avatar_find('club', cid, 'raw')
			dict_render['head_grade'] = UserSnap.grade_translate \
														(head_account.grade)
			dict_render['club'] = club
			# 评分表
			global_rubric = AuthSnap.get_rubric()
			category_rubric = AuthSnap.get_rubric(club.category)
			score = json.loads(performance.score or '[[], []]')
			dict_render['global_rubric'] = itertools.zip_longest \
													(global_rubric, score[0])
			dict_render['category_rubric'] = itertools.zip_longest \
													(category_rubric, score[1])
			dict_render['global_rubric_len'] = len(global_rubric)
			dict_render['category_rubric_len'] = len(category_rubric)
			if len(score[0]) == len(global_rubric) :
				dict_render['global_score'] = ' - ' + str(sum(score[0]))
			if len(score[1]) == len(category_rubric) :
				dict_render['category_score'] = ' - ' + str(sum(score[1]))
			#活动
			event_qry = EventSnap.event_semester_find_by_cid('club', cid).order_by \
				('time_set')
			dict_render['event_list'] = []
			for i in event_qry :
				dict_render['event_list'].append({
					'data': i, 
					'join_num': EventSnap.signup_total(i.id), 
					'datetime': ChineseSnap.datetime_simp(i.time_set), 
				})
			if not dict_render['event_location'] and event_qry :
				dict_render['event_location'] = event_qry.order_by \
					('-time_update')[0].location
			#分享
			sem_start = EventSnap.semester_start()
			share_qry = ShareSnap.share_find_by_aid('club', cid)
			share_qry = share_qry.filter(time_create__gte=sem_start, \
				category__in=('handout', 'knowledge', 'event', 'news', \
				'record-granted', 'record-denied', 'record-droped', \
				'record-broken', 'record-waiting'))
			dict_render['share_list'] = []
			for i in share_qry :
				dict_render['share_list'].append(ShareDicts.share_show \
													(i.attach_uuid))
			#成员构成
			member_makeup = [0, 0, 0, 0, 0, 0, 0, 0]
			for i in ClubSnap.grade_list_by_cid(cid) :
				member_makeup[i] += 1
			member_sum = sum(member_makeup)
			dict_render['member_makeup'] = []
			for i in range(7) :
				name = UserSnap.grade_translate(i)
				style = ' progress-bar-'
				style += ('danger', 'warning', 'info', 'success', 
						'info', 'warning', 'danger', 'success')[i % 8]
				if i > 3 :
					style += ' progress-bar-striped'
				if member_makeup[i] :
					dict_render['member_makeup'].append({
						'id': i, 
						'name': name, 
						'value': member_makeup[i], 
						'percent': member_makeup[i] * 1000 // member_sum / 10, 
						'style': style, 
					})
			#星级
			
			star_exchange = (0, 40, 50, 60, 70, 80)
			
			star = AuthSnap.star_calc(cid, event_qry, member_makeup, share_qry)
			if not event_qry or not share_qry :
				star = 0
			score = AuthSnap.star_calc \
				(cid, event_qry, member_makeup, share_qry, True)
			if performance.stars < star :
				dict_render['is_extended'] = 1
			elif performance.stars == star :
				dict_render['is_available'] = 1
			else :
				dict_render['is_failed'] = 1
				dict_render['requirement'] = \
					(star_exchange[performance.stars] - score - 1) // 10 + 1
				dict_render['requirement_p1d2'] = \
					(dict_render['requirement'] + 1) // 2
				if dict_render['requirement_p1d2'] <= 0 :
					dict_render['requirement_p1d2'] = 1
			dict_render['star_tend_list'] = [0] * performance.stars + [1] * \
				(5 - performance.stars)
			dict_render['star_list'] = [0] * star + [1] * (5 - star)
			dict_render['star_tend'] = performance.stars
			fname = dict_render['club'].full_name
			dict_render['title'] = t_(' - 申请星级 - %s') % fname
			return Snap.render('empty.html', dict_render)
		# 显示网页结束
		score = request.POST.get('score')
		if score :							# 进行评分
			if score == 'category' :
				if not is_category_admin :
					raise Http403('不是联盟长')
				index = 1
				rubric = AuthSnap.get_rubric(club.category)
			elif score == 'global' :
				if not is_union :
					raise Http403('不是社联成员')
				index = 0
				rubric = AuthSnap.get_rubric()
			else :
				raise Http404('参数错误')
			if not rubric :
				raise Snap.error('请先撰写评分表')
			score = json.loads(performance.score or '[[], []]')
			data = json.loads(request.POST.get('data', '[]'))
			if len(data) != len(rubric) :
				raise Snap.error('数据错误: 长度不一致')
			number_data = []	# 将 data 中的数据类型转换为 int 或 float
			for i, j in zip(data, rubric) :
				try :
					number_data.append(j[1][j[1].index(float(i))])
				except ValueError :
					raise Snap.error(t_('“%s”数据有误') % j[0])
			score[index] = number_data
			performance.score = json.dumps(score)
			performance.save()
			return Snap.success(request, '打分成功', { 'reload': True })
		# 进行评分结束
		if not is_head :
			raise Snap.error('不是社长')
		elif request.POST.get('star', '') :	# 更改星级
			try :
				star = int(request.POST.get('star', ''))
			except ValueError :
				raise Snap.error('内容不合法')
			performance.stars=star
			performance.data = 'preparing'
			performance.save()
			return Snap.success(request, '成功更改星级')
		else :								# 提交数据
			impact = []
			for i in range(5) :
				impact.append(request.POST.get('impact%d' % i, ''))
			if not all(impact) :
				raise Snap.error('“影响力”没有填写全')
			star = AuthSnap.star_calc(cid)
			if performance.stars > star :
				raise Snap.error('星级太高')
			performance.extra = json.dumps(impact)
			performance.data = 'waiting'
			performance.save()
			PublicMessageAction.save(request, 'empty', 'new', cid, 'club', 
				ClubSnap.cid_find_by_sname('club-union'), 'club')
			return Snap.success(request, '成功提交', 
								{ 'redirect': '/auth/apply/' })

	@vary_on_cookie
	@UserAuthority.club_union_check
	def batch(request) :
		'批量下载申请表'
		dict_render = UserAgentSnap.record(request)
		if request.GET.get('type', '') != 'form' :
			raise Http404()
		# gen file name
		for i in range(10000) :
			file_name = str(uuid.uuid1())[0:8] + '.zip'
			file_path = os.path.join \
				(settings.MEDIA_ROOT, 'archives', 'batch', file_name)
			if not os.path.exists(file_path) :
				break
		# pack
		with ZipFile(file_path, 'w') as arc:
			for i in ClubSnap.club_all_find() :
				doc = AuthSnap.form_fetch(i.id)
				if doc :
					arc.write(doc, arcname='[%03d][%s].%s' % \
						(i.id, i.full_name, doc.split('.')[-1]))
		# return
		return Snap.file(request, file_path, '社团申请表.zip')

	@vary_on_cookie
	def funds_apply(request, sname, apply=False) :
		'资金申请: 添加条目'
		dict_render = UserAgentSnap.record(request)
		if apply :
			aid = ClubSnap.cid_find_by_sname(sname)
			UserAuthority.assert_permission(request, 'club', aid)
		else :
			funds_id = int(sname)
			qry = AuthInfo.objects.filter(relation='funds', stars=funds_id)
			if not qry or qry[0].data == 'deleted' :
				raise Http404()
			aid = qry[0].account_id
			UserAuthority.assert_permission(request, 'club', aid)
		dict_render['funds_form'] = AuthDicts.funds_form()
		if request.method == 'POST' :
			form_title = request.POST.get('title')
			if not form_title :
				raise Snap.error('忘记填写标题？')
			form_type = request.POST.get('type')
			if form_type not in ('funds', 'borrow', ) :
				raise Snap.error('请选择资金类型')
			combined_data = [form_title, form_type, ]
			for i in dict_render['funds_form'] :
				content = request.POST.get('content' + str(i[0]))
				if not content :
					raise Snap.error('表单填写不全，请不要留空任何一项')
				if i[0] == 1 :
					try :
						combined_data.append(int(content))
					except Exception :
						raise Snap.error('金额必须是整数')
				else :
					combined_data.append(content)
			if not apply :
				qry[0].data = 'submitted'
				qry[0].extra = json.dumps(combined_data)
				qry[0].save()
				return Snap.success(request, '修改成功', { 'redirect': 
					'/auth/funds/%s/list/' % ClubSnap.sname_find_by_cid(aid) })
			# if apply :
			for i in range(1000) :
				prefix = int(datetime.now().strftime('%y%m%d')) * 1000
				current_qry = AuthInfo.objects.filter(relation='funds', 
								stars__gte=prefix).order_by('-stars')
				if not current_qry :
					stars = prefix + 1
				else :
					stars = current_qry[0].stars + 1
				if current_qry.filter(stars=stars) :
					continue
				qry = AuthInfo(account_id=aid, relation='funds', stars=stars, 
						data='submitted', extra=json.dumps(combined_data))
				qry.save()
				return Snap.success(request, '提交成功', { 'redirect': 
					'/auth/funds/%s/list/' % ClubSnap.sname_find_by_cid(aid) })
				break
			raise Snap.error('数据错误')
		if apply :
			dict_render['choose_funds'] = True
		else :
			dict_render['funds_id'] = funds_id
			extra = json.loads(qry[0].extra)
			dict_render['extra'] = extra
			dict_render['choose_' + extra[1]] = True
			funds_form = []
			index = 2
			for i in dict_render['funds_form'] :
				funds_form.append(i + (extra[index], ))
				index += 1
			dict_render['funds_form'] = funds_form
		head = ClubSnap.head_find_by_cid(aid)
		dict_render['head'] = head
		head_account = UserSnap.account_find_by_uid(head.id)
		dict_render['head_phone'] = head_account.phone
		dict_render['avatar'] = AvatarSnap.avatar_find('club', aid, 'raw')
		dict_render['head_grade'] = UserSnap.grade_translate \
													(head_account.grade)
		dict_render['club'] = ClubSnap.club_find_by_cid(aid)
		dict_render['title'] = ' - 资金申请'
		return Snap.render('funds_apply.html', dict_render)

	@vary_on_cookie
	def funds_delete(request, fid) :
		'资金申请: 删除条目'
		dict_render = UserAgentSnap.record(request)
		funds_id = int(fid)
		qry = AuthInfo.objects.filter(relation='funds', stars=funds_id)
		if not qry or qry[0].data == 'deleted' :
			raise Http404()
		aid = qry[0].account_id
		UserAuthority.assert_permission(request, 'club', aid)
		if request.method == 'POST' :
			AuthInfo.objects.filter(id=qry[0].id).update(data='deleted')
			return Snap.success(request, '删除成功', { 'redirect': 
				'/auth/funds/%s/list/' % ClubSnap.sname_find_by_cid(aid) })
		raise Http403('此页面只接收POST请求')

	@vary_on_cookie
	def funds_download(request, fid) :
		'资金申请: 下载为 docx'
		dict_render = UserAgentSnap.record(request)
		qry = AuthInfo.objects.filter(relation='funds', stars=fid)
		if not qry or qry[0].data == 'deleted' :
			raise Http404()
		aid = qry[0].account_id
		if not aid :
			raise Http403('数据错误')
		is_union = AccountSnap.aid_verify \
			(request, 'club', ClubSnap.cid_find_by_sname('club-union'), True, True)
		is_head = AccountSnap.aid_verify(request, 'club', aid)
		if not is_head and not is_union :
			raise Http403('您没有权限查看此页面')
		extra = json.loads(qry[0].extra)
		club = ClubSnap.club_find_by_cid(aid)
		head = ClubSnap.head_find_by_cid(aid)
		real_name = '%sforms/funds/%d.docx' % (settings.MEDIA_ROOT, 
												qry[0].stars)
		raw_name = '%d_%s_%s.docx' % (qry[0].stars, club.full_name, extra[0])
		AuthSnap.funds_document_generate \
			(club, head, extra, real_name, raw_name, qry[0])
		return Snap.file(request, real_name, raw_name)

	@vary_on_cookie
	def funds_show(request, fid) :
		'资金申请: 显示内容和调整状态'
		dict_render = UserAgentSnap.record(request)
		qry = AuthInfo.objects.filter(relation='funds', stars=fid)
		if not qry or qry[0].data == 'deleted' :
			raise Http404()
		aid = qry[0].account_id
		if not aid :
			raise Http403('数据错误')
		is_union = AccountSnap.aid_verify \
			(request, 'club', ClubSnap.cid_find_by_sname('club-union'), True, True)
		is_head = AccountSnap.aid_verify(request, 'club', aid)
		if not is_head and not is_union :
			raise Http403('您没有权限查看此页面')
		funds_form = AuthDicts.funds_form()	# 和 dict_render['funds_form'] 不同
		extra = json.loads(qry[0].extra)
		if request.method == 'POST' :
			if not is_union :
				raise Snap.error('您没有权限更改状态')
			data = request.POST.get('data', '')
			if data not in ('granted', 'rejected', 'droped', ) :
				raise Snap.error('参数错误')
			union_comment = request.POST.get('union_comment', '')
			if not union_comment :
				raise Snap.error('请输入反馈意见')
			if len(funds_form) + 2 < len(extra) :	# 有评论
				extra[len(funds_form) + 2] = union_comment
			else :
				extra.append(union_comment)
			qry.update(data=data, extra=json.dumps(extra))
			return Snap.success(request, '修改成功', { 'reload': True })
		dict_render['data'] = qry[0]
		dict_render['extra'] = extra
		status_translate = AuthDicts.funds_status_translate()
		dict_render['status'] = status_translate[qry[0].data]
		dict_render['funds_type'] = {'funds': '社团发展基金', 'borrow': '社团借款',
								}[extra[1]]
		dict_render['funds_form'] = []
		for i in range(len(funds_form)) :
			dict_render['funds_form'].append(funds_form[i] + (extra[i + 2], ))
		if len(funds_form) + 2 < len(extra) :	# 有评论
			dict_render['union_comment'] = extra[len(funds_form) + 2]
		dict_render['is_head'] = is_head
		dict_render['is_union'] = is_union
		head = ClubSnap.head_find_by_cid(aid)
		dict_render['head'] = head
		head_account = UserSnap.account_find_by_uid(head.id)
		dict_render['head_phone'] = head_account.phone
		dict_render['avatar'] = AvatarSnap.avatar_find('club', aid, 'raw')
		dict_render['head_grade'] = UserSnap.grade_translate \
															(head_account.grade)
		dict_render['club'] = ClubSnap.club_find_by_cid(aid)
		fname = dict_render['club'].full_name
		dict_render['title'] = t_(' - 资金申请状态 - %s') % fname
		return Snap.render('funds_show.html', dict_render)

	@vary_on_cookie
	@UserAuthority.club_union_check
	def funds_all(request) :
		'''
			资金申请: 显示所有
			?download=true	下载压缩包
		'''
		dict_render = UserAgentSnap.record(request)
		is_union = AccountSnap.aid_verify \
			(request, 'club', ClubSnap.cid_find_by_sname('club-union'), True, True)
		qry = AuthInfo.objects.filter(relation='funds').order_by('-time_update')
		if request.GET.get('download') == 'true' :
			name_list = []
			for i in qry.filter(data='granted') :
				club = ClubSnap.club_find_by_cid(i.account_id)
				head = ClubSnap.head_find_by_cid(i.account_id)
				extra = json.loads(i.extra)
				real_name = '%sforms/funds/%d.docx' % (settings.MEDIA_ROOT, 
														i.stars)
				raw_name = '%d_%s_%s.docx' % (i.stars, club.full_name
														, extra[0])
				name_list.append((real_name, raw_name))
				AuthSnap.funds_document_generate \
						(club, head, extra, real_name, raw_name, i)
			# gen file name
			for i in range(10000) :
				file_name = str(uuid.uuid1())[0:8] + '.zip'
				file_path = os.path.join \
					(settings.MEDIA_ROOT, 'archives', 'funds', file_name)
				if not os.path.exists(file_path) :
					break
			# pack
			with ZipFile(file_path, 'w') as arc:
				for i in name_list :
					arc.write(i[0], arcname=i[1])
			# return
			return Snap.file(request, file_path, '社团资金申请表.zip')
		dict_render['funds_list'] = AuthDicts.funds_list(qry)
		dict_render['title'] = ' - 资金申请列表'
		return Snap.render('funds_all.html', dict_render)

	@vary_on_cookie
	def funds_list(request, sname) :
		'资金申请: 显示社团的所有'
		dict_render = UserAgentSnap.record(request)
		aid = ClubSnap.cid_find_by_sname(sname)
		UserAuthority.assert_permission(request, 'club', aid)
		head = ClubSnap.head_find_by_cid(aid)
		dict_render['head'] = head
		head_account = UserSnap.account_find_by_uid(head.id)
		dict_render['head_phone'] = head_account.phone
		dict_render['avatar'] = AvatarSnap.avatar_find('club', aid, 'raw')
		dict_render['head_grade'] = UserSnap.grade_translate \
														(head_account.grade)
		dict_render['club'] = ClubSnap.club_find_by_cid(aid)
		qry = AuthInfo.objects.filter(account_id=aid, relation='funds'
							).exclude(data='deleted').order_by('-time_update')
		dict_render['funds_list'] = AuthDicts.funds_list(qry, need_club=False)
		fname = dict_render['club'].full_name
		dict_render['title'] = t_(' - 资金申请列表 - %s') % fname
		return Snap.render('funds_list.html', dict_render)

